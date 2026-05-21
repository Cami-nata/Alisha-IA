import os
import subprocess
import time
import webbrowser
import shutil
from datetime import datetime
from pathlib import Path

import pyautogui
from docx import Document

# ── NaturalMouse (fail-silent) ────────────────────────────────────────────────
try:
    from natural_mouse import NaturalMouse as _NaturalMouse
    _natural_mouse = _NaturalMouse()
    _NATURAL_MOUSE_OK = True
except Exception:
    _natural_mouse = None
    _NATURAL_MOUSE_OK = False

try:
    import psutil  # type: ignore[import]
except ImportError:
    psutil = None

# ── win32gui para foco de ventana ─────────────────────────────────────────────
try:
    import win32gui
    import win32con
    import win32process
    _WIN32_OK = True
except ImportError:
    _WIN32_OK = False

from config import APP_RUTAS, ALLOWED_APPS, POWER_COMMANDS
import app_discovery

pyautogui.FAILSAFE = True
pyautogui.PAUSE = 0.3   # reducido para que las acciones sean más fluidas


# ── Foco de ventana — bypass completo para Windows 11 ────────────────────────

def enfocar_ventana(nombre_parcial: str, abrir_si_no_existe: str = None) -> bool:
    """
    Busca una ventana por nombre parcial y le da el foco usando múltiples métodos.

    Windows 11 bloquea SetForegroundWindow si el proceso no tiene el foco.
    Solución: truco ALT + múltiples fallbacks + click físico como último recurso.

    Retorna True si logró enfocar, False si no.
    """
    nombre_lower = nombre_parcial.lower()

    # ── Paso 1: Buscar el hwnd ────────────────────────────────────────────────
    hwnd = _buscar_hwnd(nombre_lower)

    if not hwnd:
        # No encontrada — abrir si se especificó
        if abrir_si_no_existe:
            try:
                abrir_app(abrir_si_no_existe)
                time.sleep(2.5)
                hwnd = _buscar_hwnd(nombre_lower)
            except Exception as e:
                print(f"[Actions] Error abriendo '{abrir_si_no_existe}': {e}")

    if not hwnd:
        return False

    # ── Paso 2: Intentar enfocar con 3 métodos en cascada ────────────────────
    for intento in range(3):
        if _intentar_foco(hwnd):
            time.sleep(0.3)
            return True
        time.sleep(0.2)

    # ── Paso 3: Fallback final — click físico en el centro de la ventana ─────
    return _foco_por_click(hwnd)


def _buscar_hwnd(nombre_lower: str) -> int:
    """Busca el hwnd de una ventana por nombre parcial."""
    if not _WIN32_OK:
        return 0
    resultado = [0]
    def _cb(hwnd, _):
        if win32gui.IsWindowVisible(hwnd):
            titulo = win32gui.GetWindowText(hwnd).lower()
            if nombre_lower in titulo:
                resultado[0] = hwnd
        return True
    try:
        win32gui.EnumWindows(_cb, None)
    except Exception:
        pass
    return resultado[0]


def _intentar_foco(hwnd: int) -> bool:
    """
    Intenta dar foco a una ventana usando el truco ALT de Windows.

    Windows bloquea SetForegroundWindow si el proceso llamador no tiene el foco.
    El truco: simular ALT hace que Windows crea que hay interacción del usuario,
    lo que desbloquea el cambio de foco.
    """
    if not _WIN32_OK or not hwnd:
        return False
    try:
        import ctypes

        # Restaurar si está minimizada
        if win32gui.IsIconic(hwnd):
            win32gui.ShowWindow(hwnd, win32con.SW_RESTORE)
            time.sleep(0.2)

        # Truco ALT: simular pulsación de ALT para desbloquear SetForegroundWindow
        # INPUT_KEYBOARD = 1, KEYEVENTF_KEYUP = 2, VK_MENU = 0x12 (ALT)
        INPUT_KEYBOARD = 1
        KEYEVENTF_KEYUP = 0x0002
        VK_MENU = 0x12

        class KEYBDINPUT(ctypes.Structure):
            _fields_ = [
                ("wVk",         ctypes.c_ushort),
                ("wScan",       ctypes.c_ushort),
                ("dwFlags",     ctypes.c_ulong),
                ("time",        ctypes.c_ulong),
                ("dwExtraInfo", ctypes.POINTER(ctypes.c_ulong)),
            ]

        class INPUT(ctypes.Structure):
            class _INPUT(ctypes.Union):
                _fields_ = [("ki", KEYBDINPUT)]
            _anonymous_ = ("_input",)
            _fields_ = [("type", ctypes.c_ulong), ("_input", _INPUT)]

        # Presionar ALT
        inp_down = INPUT(type=INPUT_KEYBOARD)
        inp_down.ki.wVk = VK_MENU
        inp_down.ki.dwFlags = 0
        ctypes.windll.user32.SendInput(1, ctypes.byref(inp_down), ctypes.sizeof(INPUT))

        # SetForegroundWindow con ALT presionado
        win32gui.SetForegroundWindow(hwnd)

        # Soltar ALT
        inp_up = INPUT(type=INPUT_KEYBOARD)
        inp_up.ki.wVk = VK_MENU
        inp_up.ki.dwFlags = KEYEVENTF_KEYUP
        ctypes.windll.user32.SendInput(1, ctypes.byref(inp_up), ctypes.sizeof(INPUT))

        # Verificar que realmente tomó el foco
        time.sleep(0.15)
        foco_actual = ctypes.windll.user32.GetForegroundWindow()
        if foco_actual == hwnd:
            titulo = win32gui.GetWindowText(hwnd)
            print(f"[Actions] ✓ Foco obtenido (ALT trick): {titulo}")
            return True

        # Segundo intento: BringWindowToTop
        ctypes.windll.user32.BringWindowToTop(hwnd)
        ctypes.windll.user32.SetForegroundWindow(hwnd)
        time.sleep(0.15)
        foco_actual = ctypes.windll.user32.GetForegroundWindow()
        if foco_actual == hwnd:
            print(f"[Actions] ✓ Foco obtenido (BringToTop): {win32gui.GetWindowText(hwnd)}")
            return True

        return False

    except Exception as e:
        print(f"[Actions] Error en _intentar_foco: {e}")
        return False


def _foco_por_click(hwnd: int) -> bool:
    """
    Último recurso: mueve el mouse al centro de la ventana y hace click.
    Esto siempre funciona porque es una acción física del usuario.
    """
    if not _WIN32_OK or not hwnd:
        return False
    try:
        rect = win32gui.GetWindowRect(hwnd)
        x = (rect[0] + rect[2]) // 2
        y = (rect[1] + rect[3]) // 2
        # Usar NaturalMouse si está disponible
        if _NATURAL_MOUSE_OK and _natural_mouse:
            _natural_mouse.click_natural(x, y)
        else:
            pyautogui.moveTo(x, y, duration=0.4)
            pyautogui.click()
        time.sleep(0.3)
        print(f"[Actions] ✓ Foco por click físico en ({x}, {y})")
        return True
    except Exception as e:
        print(f"[Actions] Error en _foco_por_click: {e}")
        return False


def _detectar_app_destino(texto_usuario: str = "") -> tuple[str, str]:
    """
    Detecta qué app tiene el foco actualmente o cuál debería tener.
    Retorna (nombre_ventana_buscar, nombre_app_abrir).
    """
    texto_lower = texto_usuario.lower()

    # Mapeo de palabras clave → (ventana a buscar, app a abrir si no existe)
    APPS_MAPA = {
        "bloc de notas": ("bloc de notas", "notepad"),
        "notepad":       ("bloc de notas", "notepad"),
        "word":          ("word", "word"),
        "excel":         ("excel", "excel"),
        "powerpoint":    ("powerpoint", "powerpoint"),
        "vscode":        ("visual studio code", "vscode"),
        "vs code":       ("visual studio code", "vscode"),
        "code":          ("visual studio code", "vscode"),
        "chrome":        ("chrome", "chrome"),
        "edge":          ("edge", "edge"),
        "notepad++":     ("notepad++", "notepad++"),
    }

    for kw, (ventana, app) in APPS_MAPA.items():
        if kw in texto_lower:
            return ventana, app

    # Sin app específica — usar la ventana activa que no sea el navegador/chat
    return "", ""


def abrir_app(app):
    if not app:
        raise ValueError("No se indicó ninguna aplicación a abrir.")

    app_name = str(app).strip()
    llave = app_name.lower()

    # Caso especial: explorador de Windows
    if llave == "explorador":
        subprocess.Popen(
            ["cmd", "/c", "start", "", "explorer"],
            creationflags=subprocess.CREATE_NO_WINDOW,
        )
        return

    # Intentar resolver con app_discovery
    try:
        ruta = app_discovery.resolver_app(llave)
        subprocess.Popen(
            [ruta],
            creationflags=subprocess.CREATE_NO_WINDOW,
        )
        return
    except ValueError:
        pass

    raise ValueError(
        f"Aplicación no encontrada: {app_name}. "
        "Verifica que esté instalada en el sistema."
    )


def abrir_web(url):
    webbrowser.open(url)


def escribir_texto(texto, ventana_destino: str = ""):
    """
    Escribe texto en la ventana activa o en la ventana especificada.
    Usa click físico para garantizar el foco — más confiable que SetForegroundWindow.
    """
    if texto is None:
        return

    if ventana_destino:
        ventana_lower = ventana_destino.lower()
        APP_FALLBACK = {
            "bloc de notas": "notepad",
            "notepad":       "notepad",
            "word":          "word",
            "excel":         "excel",
            "powerpoint":    "powerpoint",
        }
        app_abrir = APP_FALLBACK.get(ventana_lower, ventana_lower)

        # Buscar hwnd y hacer click físico en el centro de la ventana
        hwnd = _buscar_hwnd(ventana_lower)
        if not hwnd:
            try:
                abrir_app(app_abrir)
                time.sleep(2.5)
                hwnd = _buscar_hwnd(ventana_lower)
            except Exception:
                pass

        if hwnd:
            # Intentar foco con ALT trick primero, luego click físico
            if not _intentar_foco(hwnd):
                _foco_por_click(hwnd)
        else:
            print(f"[Actions] ⚠ No se encontró '{ventana_destino}'")

    time.sleep(0.3)

    # Escribir — pyperclip+Ctrl+V es más confiable para texto con tildes/ñ
    try:
        try:
            import pyperclip
            pyperclip.copy(str(texto))
            pyautogui.hotkey("ctrl", "v")
            time.sleep(0.2)
            return
        except Exception:
            pass
        pyautogui.write(str(texto), interval=0.04)
    except pyautogui.FailSafeException:
        print("[Actions] FailSafe activado — escritura cancelada")
    except Exception as e:
        print(f"[Actions] Error escribiendo texto: {e}")


def hotkey(teclas):
    if not teclas:
        return
    pyautogui.hotkey(*teclas)


def click_xy(x, y):
    """Haz clic en las coordenadas absolutas de la pantalla.

    El origen (0,0) es la esquina superior izquierda.
    Usa NaturalMouse si está disponible; si no, pyautogui directamente.
    """
    try:
        x = int(x)
        y = int(y)
    except (TypeError, ValueError):
        raise ValueError("Coordenadas x y y deben ser números enteros.")
    if _NATURAL_MOUSE_OK and _natural_mouse is not None:
        _natural_mouse.click_natural(x, y)
    else:
        pyautogui.moveTo(x, y, duration=0.18)
        pyautogui.click(x=x, y=y)


def doble_click_xy(x, y):
    try:
        x = int(x)
        y = int(y)
    except (TypeError, ValueError):
        raise ValueError("Coordenadas x y y deben ser números enteros.")
    if _NATURAL_MOUSE_OK and _natural_mouse is not None:
        _natural_mouse.mover_a(x, y)
    else:
        pyautogui.moveTo(x, y, duration=0.18)
    pyautogui.doubleClick(x=x, y=y)


def screenshot(nombre: str = None):
    carpeta = Path(__file__).parent / "static" / "tmp"
    carpeta.mkdir(parents=True, exist_ok=True)
    if nombre:
        ruta = carpeta / nombre
    else:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        ruta = carpeta / f"captura_{timestamp}.png"
    pyautogui.screenshot(str(ruta))
    return str(ruta)


def crear_word(archivo="trabajo.docx", texto="Documento creado por la IA."):
    if not archivo.lower().endswith(".docx"):
        archivo += ".docx"
    doc = Document()
    doc.add_heading("Trabajo automático", 0)
    doc.add_paragraph(str(texto))
    doc.save(archivo)
    return archivo


def tomar_nota(titulo="nota", texto=""):
    titulo_base = str(titulo).strip() or "nota"
    titulo_seguro = "".join(ch for ch in titulo_base if ch.isalnum() or ch in " _-()").strip()
    if not titulo_seguro:
        titulo_seguro = "nota"
    nombre_archivo = f"{titulo_seguro}.txt"
    with open(nombre_archivo, "a", encoding="utf-8") as f:
        f.write(str(texto).strip() + "\n")
    return nombre_archivo


def diagnosticar_pc():
    try:
        disco = shutil.disk_usage('.')
        memoria_total = psutil.virtual_memory().total if psutil is not None else None
    except Exception:
        disco = None
        memoria_total = None

    diagnostico = ["Diagnóstico rápido de sistema:"]
    if disco:
        diagnostico.append(f"Espacio total disco: {disco.total // (1024**3)} GB")
        diagnostico.append(f"Espacio libre disco: {disco.free // (1024**3)} GB")
    else:
        diagnostico.append("No se pudo obtener información de disco.")
    if memoria_total:
        diagnostico.append(f"Memoria total: {memoria_total // (1024**3)} GB")
    else:
        diagnostico.append("No se pudo obtener información de memoria.")
    return "\n".join(diagnostico)


def control_ventana(subaccion):
    if subaccion == "minimizar":
        pyautogui.hotkey("win", "down")
    elif subaccion == "maximizar":
        pyautogui.hotkey("win", "up")
    elif subaccion == "restaurar":
        pyautogui.hotkey("win", "down")
    elif subaccion == "cerrar":
        pyautogui.hotkey("alt", "f4")
    elif subaccion == "alternar":
        pyautogui.hotkey("alt", "tab")
    elif subaccion == "mostrar_escritorio":
        pyautogui.hotkey("win", "d")


def power_action(subaccion):
    comando = POWER_COMMANDS.get(subaccion)
    if comando:
        subprocess.Popen(
            comando.split(),
            creationflags=subprocess.CREATE_NO_WINDOW,
        )