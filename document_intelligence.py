"""
document_intelligence.py — Procesador avanzado de documentos para Alisha.

Extiende file_analyzer.py con:
  - Análisis crítico de ortografía y redacción
  - Detección de contradicciones entre secciones
  - Procesamiento asíncrono para archivos pesados (>10MB)
  - Caché inteligente para archivos frecuentes
  - Integración con brain.py (SarcasmScoreEngine + MicroGestureEngine)
  - BufferMonitor para acceso en tiempo real al editor

Uso:
    from document_intelligence import DocumentIntelligence
    di = DocumentIntelligence()
    result = di.analyze("informe.docx", pregunta="¿Hay errores?")
"""

from __future__ import annotations

import hashlib
import json
import re
import threading
import time
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Callable, Dict, List, Optional

# ── Dependencias opcionales ────────────────────────────────────────────────────
try:
    import fitz  # PyMuPDF
    _FITZ_OK = True
except ImportError:
    _FITZ_OK = False

try:
    from docx import Document as _DocxDocument
    _DOCX_OK = True
except ImportError:
    _DOCX_OK = False

# ── Importar módulos existentes del proyecto ──────────────────────────────────
from file_analyzer import analizar_documento, analizar_imagen, es_archivo_soportado


# ══════════════════════════════════════════════════════════════════════════════
# DATA CLASSES
# ══════════════════════════════════════════════════════════════════════════════

@dataclass
class DocumentError:
    tipo:        str    # "ortografia" | "inconsistencia" | "redaccion" | "estructura"
    descripcion: str
    ubicacion:   str    # "sección X" o "párrafo Y"
    severidad:   float  # 0.0-1.0


@dataclass
class DocumentAnalysis:
    content:     str
    errors:      List[DocumentError] = field(default_factory=list)
    suggestions: List[str]           = field(default_factory=list)
    sarcasm_score: float             = 0.0
    word_count:  int                 = 0
    sections:    List[str]           = field(default_factory=list)
    processing_time: float           = 0.0
    from_cache:  bool                = False


# ══════════════════════════════════════════════════════════════════════════════
# CACHÉ INTELIGENTE
# ══════════════════════════════════════════════════════════════════════════════

class DocumentCache:
    """Caché en memoria para archivos frecuentemente accedidos."""

    MAX_ENTRIES = 20
    TTL_SECONDS = 3600  # 1 hora

    def __init__(self):
        self._cache: Dict[str, dict] = {}
        self._lock = threading.Lock()

    def _key(self, path: str) -> str:
        """Hash del path + mtime para invalidar si el archivo cambia."""
        try:
            mtime = Path(path).stat().st_mtime
            return hashlib.md5(f"{path}:{mtime}".encode()).hexdigest()
        except Exception:
            return hashlib.md5(path.encode()).hexdigest()

    def get(self, path: str) -> Optional[DocumentAnalysis]:
        key = self._key(path)
        with self._lock:
            entry = self._cache.get(key)
            if entry and (time.time() - entry["ts"]) < self.TTL_SECONDS:
                result = entry["data"]
                result.from_cache = True
                return result
        return None

    def set(self, path: str, analysis: DocumentAnalysis) -> None:
        key = self._key(path)
        with self._lock:
            # Evict oldest if full
            if len(self._cache) >= self.MAX_ENTRIES:
                oldest = min(self._cache, key=lambda k: self._cache[k]["ts"])
                del self._cache[oldest]
            self._cache[key] = {"data": analysis, "ts": time.time()}

    def invalidate(self, path: str) -> None:
        key = self._key(path)
        with self._lock:
            self._cache.pop(key, None)


# ══════════════════════════════════════════════════════════════════════════════
# CRITIC ENGINE — análisis crítico de texto
# ══════════════════════════════════════════════════════════════════════════════

class CriticEngine:
    """
    Analiza texto en busca de errores ortográficos, inconsistencias
    lógicas y problemas de redacción.
    """

    # Errores ortográficos comunes en español rioplatense
    _ORTOGRAFIA_COMUN = [
        (r'\baber\b',        'haber',       'ortografía'),
        (r'\balla\b',        'haya/allá',   'ortografía'),
        (r'\bvalla\b',       'vaya/valla',  'ortografía'),
        (r'\becho\b(?!\s+(?:un|una|el|la|los|las|de|a|en))', 'hecho', 'ortografía'),
        (r'\basi\b',         'así',         'tilde'),
        (r'\bmas\b(?!\s+\w+\s+que)', 'más', 'tilde'),
        (r'\bsolo\b',        'sólo/solo',   'tilde'),
        (r'\bque\b(?=\s+[A-Z])', 'qué',    'tilde'),
        (r'\bcomo\b(?=\s+[A-Z])', 'cómo',  'tilde'),
        (r'\bporque\b(?=\s*\?)', 'por qué', 'ortografía'),
        (r'\bpor que\b(?!\s*\?)', 'porque/por qué', 'ortografía'),
        (r'\ba ver\b',       'a ver',       'ok'),  # correcto
        (r'\baver\b',        'a ver',       'ortografía'),
    ]

    # Patrones de inconsistencia lógica
    _INCONSISTENCIAS = [
        (r'(es|son|será|serán)\s+\w+.*?(no es|no son|nunca es|nunca son)',
         'posible contradicción afirmación/negación'),
        (r'(siempre|nunca|todos|ninguno).*?(a veces|algunos|pocos)',
         'posible contradicción absoluto/relativo'),
        (r'(aumenta|crece|sube).*?(disminuye|baja|cae)',
         'posible contradicción de tendencia'),
        (r'(objetivo|meta|fin)\s+es.*?(objetivo|meta|fin)\s+es',
         'objetivos múltiples sin jerarquía'),
    ]

    def analyze(self, text: str) -> List[DocumentError]:
        errors = []
        errors.extend(self._check_ortografia(text))
        errors.extend(self._check_inconsistencias(text))
        errors.extend(self._check_fechas(text))
        errors.extend(self._check_redaccion(text))
        return errors

    def _check_fechas(self, text: str) -> List[DocumentError]:
        """
        Detecta contradicciones de fechas/años en el documento.
        Ej: menciona 2026 al principio y 2024 al final.
        """
        errors = []
        # Buscar todos los años de 4 dígitos (1900-2099)
        years = re.findall(r'\b(19\d{2}|20\d{2})\b', text)
        if len(years) >= 2:
            unique_years = list(dict.fromkeys(years))  # preservar orden, sin duplicados
            if len(unique_years) >= 2:
                # Verificar si hay años distintos en primera y segunda mitad del doc
                mid = len(text) // 2
                years_first  = set(re.findall(r'\b(19\d{2}|20\d{2})\b', text[:mid]))
                years_second = set(re.findall(r'\b(19\d{2}|20\d{2})\b', text[mid:]))
                conflicting  = years_first & years_second
                # Si hay años distintos entre primera y segunda mitad
                diff_first  = years_first  - years_second
                diff_second = years_second - years_first
                if diff_first and diff_second:
                    errors.append(DocumentError(
                        tipo='inconsistencia',
                        descripcion=(
                            f"contradicción de fechas: "
                            f"primera mitad menciona {sorted(diff_first)} "
                            f"y segunda mitad menciona {sorted(diff_second)}"
                        ),
                        ubicacion='documento completo',
                        severidad=0.85,
                    ))
        return errors

    def _check_ortografia(self, text: str) -> List[DocumentError]:
        errors = []
        for pattern, correcto, tipo in self._ORTOGRAFIA_COMUN:
            if tipo == 'ok':
                continue
            matches = re.finditer(pattern, text, re.IGNORECASE)
            for m in matches:
                errors.append(DocumentError(
                    tipo='ortografia',
                    descripcion=f'"{m.group()}" → debería ser "{correcto}"',
                    ubicacion=f'posición {m.start()}',
                    severidad=0.3 if tipo == 'tilde' else 0.5,
                ))
        return errors

    def _check_inconsistencias(self, text: str) -> List[DocumentError]:
        errors = []
        text_lower = text.lower()
        for pattern, desc in self._INCONSISTENCIAS:
            if re.search(pattern, text_lower):
                errors.append(DocumentError(
                    tipo='inconsistencia',
                    descripcion=desc,
                    ubicacion='documento completo',
                    severidad=0.7,
                ))
        return errors

    def _check_redaccion(self, text: str) -> List[DocumentError]:
        errors = []
        paragraphs = [p.strip() for p in text.split('\n') if p.strip()]

        for i, para in enumerate(paragraphs):
            # Párrafos muy cortos (menos de 20 chars) que no son títulos
            if 5 < len(para) < 20 and not para.endswith(':'):
                errors.append(DocumentError(
                    tipo='redaccion',
                    descripcion=f'Párrafo muy corto: "{para[:30]}"',
                    ubicacion=f'párrafo {i+1}',
                    severidad=0.2,
                ))

            # Oraciones muy largas (más de 200 chars sin punto)
            if len(para) > 200 and para.count('.') == 0:
                errors.append(DocumentError(
                    tipo='redaccion',
                    descripcion='Oración muy larga sin puntuación',
                    ubicacion=f'párrafo {i+1}',
                    severidad=0.3,
                ))

        return errors

    def calculate_sarcasm_score(self, errors: List[DocumentError]) -> float:
        """Calcula el Sarcasm Score basado en errores detectados."""
        score = 0.0
        for err in errors:
            if err.tipo == 'ortografia':
                score += 0.1
            elif err.tipo == 'inconsistencia':
                score += 0.25
            elif err.tipo == 'redaccion':
                score += 0.05
        return min(score, 1.0)

    def generate_suggestions(self, errors: List[DocumentError]) -> List[str]:
        """Genera sugerencias de corrección en tono rioplatense."""
        if not errors:
            return ["Todo bien por acá, no encontré errores graves."]

        suggestions = []
        ortografia = [e for e in errors if e.tipo == 'ortografia']
        inconsistencias = [e for e in errors if e.tipo == 'inconsistencia']
        redaccion = [e for e in errors if e.tipo == 'redaccion']

        if ortografia:
            suggestions.append(
                f"Encontré {len(ortografia)} error(es) ortográfico(s). "
                f"Ejemplo: {ortografia[0].descripcion}"
            )
        if inconsistencias:
            suggestions.append(
                f"Hay {len(inconsistencias)} posible(s) contradicción(es) en el texto. "
                "Revisá que no te estés contradiciendo entre secciones."
            )
        if redaccion:
            suggestions.append(
                f"La redacción tiene {len(redaccion)} punto(s) a mejorar. "
                "Algunos párrafos podrían estar mejor estructurados."
            )

        return suggestions


# ══════════════════════════════════════════════════════════════════════════════
# BUFFER MONITOR — acceso en tiempo real al editor
# ══════════════════════════════════════════════════════════════════════════════

class BufferMonitor:
    """
    Monitorea el buffer del editor activo en tiempo real.
    Notifica a Alisha cuando el contenido cambia.
    """

    POLL_INTERVAL = 2.0  # segundos entre chequeos

    def __init__(self):
        self._last_content: str = ""
        self._last_hash: str = ""
        self._running = False
        self._thread: Optional[threading.Thread] = None
        self._callbacks: List[Callable[[str], None]] = []

    def on_change(self, callback: Callable[[str], None]) -> None:
        """Registra un callback que se llama cuando el buffer cambia."""
        self._callbacks.append(callback)

    def start(self) -> None:
        self._running = True
        self._thread = threading.Thread(target=self._monitor, daemon=True)
        self._thread.start()
        print("[BufferMonitor] ✓ Iniciado")

    def stop(self) -> None:
        self._running = False

    def get_current_content(self) -> str:
        return self._last_content

    def _monitor(self) -> None:
        while self._running:
            try:
                content = self._read_active_editor()
                if content:
                    content_hash = hashlib.md5(content.encode()).hexdigest()
                    if content_hash != self._last_hash:
                        self._last_hash = content_hash
                        self._last_content = content
                        for cb in self._callbacks:
                            try:
                                cb(content)
                            except Exception:
                                pass
            except Exception:
                pass
            time.sleep(self.POLL_INTERVAL)

    def _read_active_editor(self) -> Optional[str]:
        """Intenta leer el contenido del editor activo."""
        # Intentar screen_vision primero (VS Code)
        try:
            from screen_vision import leer_vscode_activo
            content = leer_vscode_activo()
            if content:
                return content
        except Exception:
            pass

        # Fallback: leer desde clipboard si hay texto copiado
        try:
            import pyperclip
            clip = pyperclip.paste()
            if clip and len(clip) > 50:
                return clip
        except Exception:
            pass

        return None


# ══════════════════════════════════════════════════════════════════════════════
# DOCUMENT INTELLIGENCE — núcleo principal
# ══════════════════════════════════════════════════════════════════════════════

class DocumentIntelligence:
    """
    Procesador avanzado de documentos con análisis crítico integrado.

    Uso:
        di = DocumentIntelligence()
        result = di.analyze("informe.docx")
        print(result.suggestions)
    """

    HEAVY_FILE_THRESHOLD = 10 * 1024 * 1024  # 10MB

    def __init__(self):
        self._cache   = DocumentCache()
        self._critic  = CriticEngine()
        self._buffer  = BufferMonitor()
        self._lock    = threading.Lock()

    # ── API pública ────────────────────────────────────────────────────────────

    def analyze(self, path: str, pregunta: str = "",
                use_cache: bool = True) -> DocumentAnalysis:
        """
        Analiza un documento y retorna un DocumentAnalysis completo.
        Para archivos >10MB usa procesamiento por segmentos.
        """
        t_start = time.time()

        # Verificar caché
        if use_cache:
            cached = self._cache.get(path)
            if cached:
                print(f"[DocIntelligence] ✓ Desde caché: {Path(path).name}")
                return cached

        # Activar micro-gesto de análisis
        self._trigger_analysis_gesture()

        file_size = Path(path).stat().st_size if Path(path).exists() else 0
        is_heavy  = file_size > self.HEAVY_FILE_THRESHOLD

        if is_heavy:
            print(f"[DocIntelligence] 📦 Archivo pesado ({file_size/1024/1024:.1f}MB) — procesando por segmentos")
            content = self._extract_segmented(path)
        else:
            content = self._extract_content(path)

        if not content or content.startswith("[Error"):
            return DocumentAnalysis(
                content=content or "No se pudo extraer contenido",
                errors=[],
                suggestions=["No pude leer el archivo. ¿Está bien formado?"],
            )

        # Análisis crítico
        errors      = self._critic.analyze(content)
        suggestions = self._critic.generate_suggestions(errors)
        sarcasm     = self._critic.calculate_sarcasm_score(errors)
        sections    = self._extract_sections(content)

        # Si hay errores graves, activar modo crítico en Live2D
        if sarcasm > 0.4:
            self._trigger_critic_gesture(sarcasm)

        # Análisis con LLM si hay pregunta específica
        if pregunta:
            llm_analysis = self._analyze_with_llm(path, content, pregunta)
            suggestions.append(llm_analysis)

        result = DocumentAnalysis(
            content=content,
            errors=errors,
            suggestions=suggestions,
            sarcasm_score=sarcasm,
            word_count=len(content.split()),
            sections=sections,
            processing_time=time.time() - t_start,
        )

        # Guardar en caché
        self._cache.set(path, result)

        print(f"[DocIntelligence] ✓ Analizado en {result.processing_time:.2f}s "
              f"({result.word_count} palabras, {len(errors)} errores, sarcasmo={sarcasm:.2f})")

        return result

    def analyze_async(self, path: str, callback: Callable[[DocumentAnalysis], None],
                      pregunta: str = "") -> None:
        """Versión asíncrona — no bloquea la interfaz."""
        def _run():
            result = self.analyze(path, pregunta)
            callback(result)

        threading.Thread(target=_run, daemon=True).start()

    def analyze_text(self, text: str) -> DocumentAnalysis:
        """Analiza texto directo (sin archivo)."""
        t_start = time.time()
        errors      = self._critic.analyze(text)
        suggestions = self._critic.generate_suggestions(errors)
        sarcasm     = self._critic.calculate_sarcasm_score(errors)

        if sarcasm > 0.4:
            self._trigger_critic_gesture(sarcasm)

        return DocumentAnalysis(
            content=text,
            errors=errors,
            suggestions=suggestions,
            sarcasm_score=sarcasm,
            word_count=len(text.split()),
            processing_time=time.time() - t_start,
        )

    def get_buffer_monitor(self) -> BufferMonitor:
        return self._buffer

    def start_buffer_monitoring(self) -> None:
        self._buffer.start()

    # ── Extracción de contenido ────────────────────────────────────────────────

    def _extract_content(self, path: str) -> str:
        ext = Path(path).suffix.lower()
        if ext == '.pdf':
            return self._extract_pdf(path)
        elif ext == '.docx':
            return self._extract_docx(path)
        elif ext in {'.txt', '.md', '.csv'}:
            return self._extract_text(path)
        else:
            return f"[Formato no soportado: {ext}]"

    def _extract_pdf(self, path: str) -> str:
        if not _FITZ_OK:
            # Fallback al file_analyzer existente
            return analizar_documento(path)
        try:
            doc = fitz.open(path)
            pages = []
            for i, page in enumerate(doc):
                text = page.get_text()
                if text.strip():
                    pages.append(f"[Página {i+1}]\n{text}")
            doc.close()
            return "\n\n".join(pages)[:15000]
        except Exception as e:
            return f"[Error leyendo PDF: {e}]"

    def _describir_imagen(self, doc, img_info: tuple, n: int) -> str:
        """Obtiene descripción de una imagen via GeminiVision. Fail-silent."""
        try:
            from gemini_vision import GeminiVision
            xref = img_info[0]
            base_image = doc.extract_image(xref)
            img_bytes = base_image["image"]
            vision = GeminiVision()
            descripcion = vision.describe_image_bytes(img_bytes)
            return f"[Imagen {n}]\n{descripcion}"
        except Exception:
            return f"[Imagen {n}: no disponible]"

    def _extract_pdf_with_vision(self, path: str) -> str:
        """
        Extrae texto e imágenes de un PDF usando PyMuPDF + GeminiVision.
        Trunca el resultado a 15.000 caracteres.
        Fallback a analizar_documento() si PyMuPDF no está instalado.
        """
        if not _FITZ_OK:
            return analizar_documento(path)

        try:
            doc = fitz.open(path)
        except Exception:
            return f"[Error: archivo no encontrado o inaccesible: {path}]"

        pages = []
        for i, page in enumerate(doc):
            texto = page.get_text().strip()
            partes = []
            if texto:
                partes.append(texto)

            # Extraer imágenes si GeminiVision está disponible
            try:
                imagenes = page.get_images(full=True)
                for n, img_info in enumerate(imagenes, start=1):
                    descripcion = self._describir_imagen(doc, img_info, n)
                    partes.append(descripcion)
            except Exception:
                pass

            if partes:
                pages.append(f"[Página {i+1}]\n" + "\n".join(partes))

        doc.close()
        contenido = "\n\n".join(pages)
        return contenido[:15000]

    def _extract_docx(self, path: str) -> str:
        if not _DOCX_OK:
            return analizar_documento(path)
        try:
            doc = _DocxDocument(path)
            sections = []
            current_section = []

            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue
                # Detectar títulos (estilo Heading)
                if para.style.name.startswith('Heading'):
                    if current_section:
                        sections.append('\n'.join(current_section))
                    current_section = [f"\n## {text}"]
                else:
                    current_section.append(text)

            if current_section:
                sections.append('\n'.join(current_section))

            return '\n\n'.join(sections)[:15000]
        except Exception as e:
            return f"[Error leyendo DOCX: {e}]"

    def _extract_text(self, path: str) -> str:
        try:
            return Path(path).read_text(encoding='utf-8', errors='ignore')[:15000]
        except Exception as e:
            return f"[Error leyendo archivo: {e}]"

    def _extract_segmented(self, path: str, segment_size: int = 5000) -> str:
        """Procesa archivos pesados por segmentos para no saturar la memoria."""
        ext = Path(path).suffix.lower()
        segments = []

        if ext == '.pdf' and _FITZ_OK:
            try:
                doc = fitz.open(path)
                total = len(doc)
                print(f"[DocIntelligence] 📄 {total} páginas — procesando...")
                for i, page in enumerate(doc):
                    text = page.get_text()
                    if text.strip():
                        segments.append(f"[Pág {i+1}/{total}] {text[:segment_size]}")
                    # Progreso cada 10 páginas
                    if (i + 1) % 10 == 0:
                        print(f"[DocIntelligence] ⏳ {i+1}/{total} páginas procesadas")
                doc.close()
            except Exception as e:
                return f"[Error procesando PDF pesado: {e}]"

        elif ext == '.docx' and _DOCX_OK:
            try:
                doc = _DocxDocument(path)
                paras = [p.text for p in doc.paragraphs if p.text.strip()]
                total = len(paras)
                for i in range(0, total, 50):
                    chunk = paras[i:i+50]
                    segments.append('\n'.join(chunk))
                    if i % 200 == 0 and i > 0:
                        print(f"[DocIntelligence] ⏳ {i}/{total} párrafos procesados")
            except Exception as e:
                return f"[Error procesando DOCX pesado: {e}]"

        return '\n\n'.join(segments)[:20000]

    def _extract_sections(self, content: str) -> List[str]:
        """Extrae títulos de secciones del contenido."""
        sections = []
        for line in content.split('\n'):
            line = line.strip()
            if line.startswith('##') or (len(line) < 80 and line.endswith(':')):
                sections.append(line.lstrip('#').strip())
        return sections[:20]

    def _analyze_with_llm(self, path: str, content: str, pregunta: str) -> str:
        """Usa el LLM para análisis específico con la pregunta del usuario."""
        try:
            from brain import get_brain
            brain = get_brain()
            # Pasar errores detectados para activar SarcasmScoreEngine
            error_strings = [
                f"{e.tipo}: {e.descripcion}" for e in self._critic.analyze(content)
            ]
            prompt = (
                f"Analizá este documento y respondé: {pregunta}\n\n"
                f"Contenido (primeros 3000 chars):\n{content[:3000]}"
            )
            response = brain.process(prompt, errors=error_strings)
            return response.content
        except Exception:
            return analizar_documento(path, pregunta)

    # ── Integración Live2D ─────────────────────────────────────────────────────

    @staticmethod
    def _trigger_analysis_gesture() -> None:
        """Activa gesto de análisis en Live2D."""
        try:
            from brain import MicroGestureEngine
            MicroGestureEngine._write_state("curiosidad", hablando=False)
        except Exception:
            pass

    @staticmethod
    def _trigger_critic_gesture(sarcasm_score: float) -> None:
        """Activa gesto crítico según nivel de sarcasmo."""
        try:
            from brain import MicroGestureEngine
            if sarcasm_score > 0.7:
                MicroGestureEngine._write_state("frustración", hablando=False)
                print("[DocIntelligence] 😏 Modo Sarcástico — muchos errores detectados")
            else:
                MicroGestureEngine._write_state("preocupación", hablando=False)
                print("[DocIntelligence] 🔍 Modo Programadora — analizando errores")
        except Exception:
            pass


# ══════════════════════════════════════════════════════════════════════════════
# SINGLETON
# ══════════════════════════════════════════════════════════════════════════════

_doc_intelligence: Optional[DocumentIntelligence] = None

def get_document_intelligence() -> DocumentIntelligence:
    global _doc_intelligence
    if _doc_intelligence is None:
        _doc_intelligence = DocumentIntelligence()
    return _doc_intelligence
